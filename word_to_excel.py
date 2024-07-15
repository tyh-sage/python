from docx import Document
from openpyxl import load_workbook

def insert_word_to_excel(word_path, excel_path, start_row=1, start_col=1, horizontal=True):
    # 加载Word文档
    doc = Document(word_path)
    # 加载Excel工作簿，默认是第一个工作表
    wb = load_workbook(excel_path)
    ws = wb.active

    # 根据输入方向设置遍历方式
    if horizontal:
        # 横向输入：先按段落遍历，再按行
        for paragraph in doc.paragraphs:
            for text in paragraph.text.split():
                if start_col <= ws.max_column:
                    ws.cell(row=start_row, column=start_col).value = text
                    start_col += 1
                else:
                    start_row += 1
                    start_col = 1
    else:
        # 列向输入：先按行遍历，再按段落
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            row = start_row + i
            if row <= ws.max_row:
                ws.cell(row=row, column=start_col).value = text
            else:
                break

    # 保存Excel工作簿
    wb.save(excel_path)

# 调用函数
word_path = 'C://Users//sage//Desktop//1.docx'  # Word文件路径——注意：文件要处于关闭状态
excel_path = 'C://Users//sage//Desktop//2.xlsx'  # Excel文件路径——注意：文件要处于关闭状态
# 指定起始行和列，以及输入方向（横向或列向）
insert_word_to_excel(word_path, excel_path, start_row=3, start_col=2, horizontal=False)
# 可以修改起始的行（从1开始的）、列（从1开始的）、横向(True)or列向(False)输入
